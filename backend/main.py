from fastapi import FastAPI, File, UploadFile, HTTPException, Form, BackgroundTasks
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import httpx
import PyPDF2
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
from datetime import datetime
import uuid
import redis
import json
from celery import Celery
import asyncio

app = FastAPI()

# CORS - Allow all origins for now
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Redis connection
try:
    redis_client = redis.from_url(
        os.getenv("REDIS_URL", "redis://localhost:6379"),
        decode_responses=True,
        socket_connect_timeout=5
    )
    redis_client.ping()
    print("✓ Redis connected successfully")
except Exception as e:
    print(f"⚠ Redis connection failed: {e}")
    redis_client = None

# Celery configuration
celery_app = Celery(
    "research_assistant",
    broker=os.getenv("REDIS_URL", "redis://localhost:6379"),
    backend=os.getenv("REDIS_URL", "redis://localhost:6379")
)

# Groq API configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

# Semaphore for rate limiting
api_semaphore = asyncio.Semaphore(10)

# In-memory fallback if Redis fails
memory_sessions = {}

class SessionData:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.research_pdfs = []
        self.research_texts = []
        self.analysis_data = None
        self.analysis_guide = None
        self.analysis_objectives = None
        self.research_responses = []
        self.analysis_steps = []
        self.analysis_results = []
    
    def to_dict(self):
        return {
            "session_id": self.session_id,
            "research_pdfs": self.research_pdfs,
            "research_texts": self.research_texts,
            "analysis_data": self.analysis_data.to_json() if self.analysis_data is not None else None,
            "analysis_guide": self.analysis_guide,
            "analysis_objectives": self.analysis_objectives,
            "research_responses": self.research_responses,
            "analysis_steps": self.analysis_steps,
            "analysis_results": self.analysis_results
        }
    
    @classmethod
    def from_dict(cls, data: dict):
        session = cls(data["session_id"])
        session.research_pdfs = data.get("research_pdfs", [])
        session.research_texts = data.get("research_texts", [])
        
        if data.get("analysis_data"):
            session.analysis_data = pd.read_json(data["analysis_data"])
        
        session.analysis_guide = data.get("analysis_guide")
        session.analysis_objectives = data.get("analysis_objectives")
        session.research_responses = data.get("research_responses", [])
        session.analysis_steps = data.get("analysis_steps", [])
        session.analysis_results = data.get("analysis_results", [])
        
        return session

def get_session(session_id: str) -> SessionData:
    # Try Redis first
    if redis_client:
        try:
            data = redis_client.get(f"session:{session_id}")
            if data:
                return SessionData.from_dict(json.loads(data))
        except Exception as e:
            print(f"Redis get error: {e}")
    
    # Fallback to memory
    if session_id in memory_sessions:
        return memory_sessions[session_id]
    
    # Create new session
    session = SessionData(session_id)
    memory_sessions[session_id] = session
    return session

def save_session(session: SessionData):
    # Save to memory first
    memory_sessions[session.session_id] = session
    
    # Try Redis
    if redis_client:
        try:
            redis_client.setex(
                f"session:{session.session_id}",
                7200,  # 2 hours TTL
                json.dumps(session.to_dict())
            )
        except Exception as e:
            print(f"Redis save error: {e}")

def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_docx_text(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def chunk_text(text: str, max_chars: int = 50000) -> List[str]:
    """Split text into chunks to avoid payload size limits"""
    chunks = []
    current_chunk = ""
    
    for paragraph in text.split('\n'):
        if len(current_chunk) + len(paragraph) > max_chars:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = paragraph
        else:
            current_chunk += "\n" + paragraph
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

async def call_groq(prompt: str, context: str = "", model: str = "llama-3.3-70b-versatile") -> str:
    """Call Groq API with rate limiting and chunking"""
    async with api_semaphore:
        # Limit total input size
        max_total_chars = 60000  # ~15k tokens
        
        if context:
            # Truncate context if too long
            if len(context) > 40000:
                context = context[:40000] + "\n\n[Context truncated due to length...]"
        
        full_prompt = f"{context}\n\n{prompt}" if context else prompt
        
        # Final check
        if len(full_prompt) > max_total_chars:
            full_prompt = full_prompt[:max_total_chars] + "\n\n[Truncated...]"
        
        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a professional research assistant with expertise in academic writing, data analysis, and research methodology. Provide concise, accurate responses."
                },
                {
                    "role": "user",
                    "content": full_prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": 4000  # Reduced from 8000
        }
        
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(GROQ_API_URL, headers=headers, json=payload)
                response.raise_for_status()
                data = response.json()
                return data["choices"][0]["message"]["content"]
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 413:
                # Payload too large - try with even smaller context
                if context:
                    truncated_context = context[:20000]
                    return await call_groq(prompt, truncated_context, model)
                else:
                    raise HTTPException(500, "Prompt too large even after truncation")
            raise HTTPException(500, f"Groq API error: {str(e)}")
        except Exception as e:
            raise HTTPException(500, f"API call failed: {str(e)}")

# Celery task for background PDF processing
@celery_app.task
def process_pdf_background(session_id: str, filename: str, content_base64: str):
    import base64
    content = base64.b64decode(content_base64)
    text = extract_pdf_text(content)
    
    session = get_session(session_id)
    session.research_pdfs.append({
        "name": filename,
        "size": len(content),
        "processed": True
    })
    session.research_texts.append(f"File: {filename}\n{text}")
    save_session(session)
    
    return {"status": "completed", "filename": filename}

class ResearchTaskRequest(BaseModel):
    prompt: str
    session_id: str

class ApprovalRequest(BaseModel):
    step_id: str
    approved: bool
    session_id: str

class ReportRequest(BaseModel):
    type: str
    sections: List[str]
    session_id: str

@app.post("/research/upload")
async def upload_research_pdfs(
    files: List[UploadFile] = File(...),
    session_id: str = Form(...)
):
    session = get_session(session_id)

    if not hasattr(session, "research_texts"):
        session.research_texts = []

    uploaded_files = []

    for file in files:
        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Only PDF files allowed")

        content = await file.read()
        text = extract_pdf_text(content)

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail=f"No text extracted from {file.filename}"
            )

        session.research_texts.append(text)
        uploaded_files.append({"name": file.filename})

    save_session(session)

    return {
        "files": uploaded_files,
        "session_id": session_id
    }

@app.post("/research/execute")
async def execute_research_task(request: ResearchTaskRequest):
    session = get_session(request.session_id)

    if not hasattr(session, "research_texts") or not session.research_texts:
        raise HTTPException(
            status_code=400,
            detail="No research PDFs uploaded or processed"
        )

    all_texts = "\n\n---\n\n".join(session.research_texts)

    max_context = 30000
    if len(all_texts) > max_context:
        all_texts = all_texts[:max_context] + "\n\n[Additional content truncated...]"

    context_prompt = f"""
You are a professional research assistant. You have access to the following research papers:

{all_texts}

Task: {request.prompt}

Use APA citation format. Max 800 words.
"""

    try:
        response = await call_groq(context_prompt)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    session.research_responses.append({
        "prompt": request.prompt,
        "response": response
    })
    save_session(session)

    return {"response": response}


@app.post("/analysis/upload")
async def upload_analysis_file(
    file: UploadFile = File(...),
    type: str = Form(...),
    session_id: str = Form(...)
):
    session = get_session(session_id)
    content = await file.read()
    
    if type == "data":
        if file.filename.endswith('.csv'):
            session.analysis_data = pd.read_csv(io.BytesIO(content))
        elif file.filename.endswith('.xlsx'):
            session.analysis_data = pd.read_excel(io.BytesIO(content))
        else:
            raise HTTPException(400, "Data file must be CSV or XLSX")
    
    elif type in ["guide", "objectives"]:
        if not file.filename.endswith(('.doc', '.docx')):
            raise HTTPException(400, "Guide and objectives must be DOC/DOCX files")
        text = extract_docx_text(content)
        if type == "guide":
            session.analysis_guide = text
        else:
            session.analysis_objectives = text
    
    save_session(session)
    return {"file": {"name": file.filename, "type": type}}

@app.post("/analysis/start")
async def start_analysis(session_id: str = Form(...)):
    session = get_session(session_id)
    
    if session.analysis_data is None or session.analysis_objectives is None:
        raise HTTPException(400, "Data and objectives files are required")
    
    # Prepare limited context
    data_info = f"""Dataset Overview:
- Columns: {', '.join(session.analysis_data.columns[:20])}  
- Rows: {len(session.analysis_data)}
- Sample (first 5 rows):
{session.analysis_data.head(5).to_string()}"""
    
    objectives_text = session.analysis_objectives[:5000]  # Limit objectives length
    
    planning_prompt = f"""Based on these objectives and dataset, create 5 clear analysis steps.

Objectives: {objectives_text}

{data_info}

Provide EXACTLY 5 steps. Format:
Step 1: [Title]
Description: [What to do]
Expected Output: [What will be produced]"""
    
    plan_response = await call_groq(planning_prompt)
    
    # Parse steps
    steps = []
    lines = plan_response.split('\n')
    current_step = None
    
    for line in lines:
        line_stripped = line.strip()
        if line_stripped.startswith('Step') and ':' in line_stripped:
            if current_step:
                steps.append(current_step)
            
            title_part = line_stripped.split(':', 1)[1].strip() if ':' in line_stripped else line_stripped
            current_step = {
                "id": str(uuid.uuid4()),
                "title": title_part[:100],
                "description": title_part,
                "completed": False
            }
        elif current_step and line_stripped and not line_stripped.startswith('Step'):
            current_step["description"] += " " + line_stripped
    
    if current_step:
        steps.append(current_step)
    
    if not steps:
        steps = [{
            "id": str(uuid.uuid4()),
            "title": "Data Analysis",
            "description": "Perform comprehensive analysis",
            "completed": False
        }]
    
    session.analysis_steps = steps
    save_session(session)
    
    return {"steps": steps[:1]}

@app.post("/analysis/approve")
async def approve_analysis_step(request: ApprovalRequest):
    session = get_session(request.session_id)
    
    if not request.approved:
        return {"status": "rejected"}
    
    step = next((s for s in session.analysis_steps if s["id"] == request.step_id), None)
    if not step:
        raise HTTPException(404, "Step not found")
    
    # Limited context for execution
    data_summary = f"""Dataset: {len(session.analysis_data)} rows, {len(session.analysis_data.columns)} columns
Sample:
{session.analysis_data.head(5).to_string()}"""
    
    execution_prompt = f"""{data_summary}

Execute: {step['description']}

Provide concise results (max 400 words):
1. Methodology
2. Key findings
3. Statistical values
4. Interpretation"""
    
    result = await call_groq(execution_prompt)
    
    session.analysis_results.append({
        "step_id": request.step_id,
        "summary": result
    })
    save_session(session)
    
    return {"result": result}

@app.post("/report/generate")
async def generate_report(request: ReportRequest):
    session = get_session(request.session_id)
    
    doc = docx.Document()
    
    title = doc.add_heading(f"{request.type.capitalize()} Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()
    
    if request.type == "research":
        for section in request.sections:
            doc.add_heading(section, 1)
            
            relevant = [r for r in session.research_responses if section.lower() in r['prompt'].lower()]
            if relevant:
                for resp in relevant:
                    doc.add_paragraph(resp['response'])
            else:
                doc.add_paragraph(f"[Content for {section} section]")
            
            doc.add_paragraph()
    
    else:  # analysis
        for section in request.sections:
            doc.add_heading(section, 1)
            
            if section == "Objectives":
                doc.add_paragraph(session.analysis_objectives or "[Not provided]")
            elif section in ["Methods", "Results", "Interpretation"]:
                for result in session.analysis_results:
                    doc.add_paragraph(result.get('summary', ''))
            else:
                doc.add_paragraph(f"[Content for {section}]")
            
            doc.add_paragraph()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={request.type}_report_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

@app.post("/session/clear")
async def clear_session(session_id: str = Form(...)):
    if redis_client:
        try:
            redis_client.delete(f"session:{session_id}")
        except Exception as e:
            print(f"Redis delete error: {e}")
    
    if session_id in memory_sessions:
        del memory_sessions[session_id]
    
    return {"status": "cleared"}

@app.get("/health")
async def health_check():
    redis_status = "disconnected"
    if redis_client:
        try:
            redis_client.ping()
            redis_status = "connected"
        except:
            redis_status = "error"
    
    return {
        "status": "healthy",
        "redis": redis_status,
        "groq_api": "configured" if GROQ_API_KEY else "missing"
    }

@app.get("/")
async def root():
    return {"message": "AI Research Assistant API with Groq", "version": "2.1"}

@app.api_route("/status", methods=["GET", "HEAD"])
def status():
    """Health check endpoint."""
    return {"status": "ok", "message": "Research Project API is running."}